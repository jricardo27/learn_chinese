import json
import re
import os
import argparse
from docx import Document
from docx.shared import Inches

def parse_js_data(file_path):
    """Parses the WORDS_DATA object from the JavaScript file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Remove 'const WORDS_DATA = ' and the trailing ';'
    content = content.strip()
    if content.startswith('const WORDS_DATA ='):
        content = content[len('const WORDS_DATA ='):].strip()
    # Find the first { and the last }
    start_idx = content.find('{')
    end_idx = content.rfind('}')
    content = content[start_idx:end_idx+1]
    
    # Handle non-standard JSON (keys without quotes, trailing commas)
    # This is a bit risky but usually works for simple objects
    content = re.sub(r'(\w+):', r'"\1":', content)
    content = re.sub(r',\s*}', '}', content)
    content = re.sub(r',\s*]', ']', content)
    
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        print(f"Fatal: Could not parse the JavaScript data file. Error: {e}")
        raise

def has_sublist(full_list, sub_list):
    """Checks if sub_list appears as a contiguous sequence in full_list."""
    n = len(full_list)
    m = len(sub_list)
    if m > n:
        return False
    for i in range(n - m + 1):
        if full_list[i:i+m] == sub_list:
            return True
    return False

def filter_words(data, tones=None, categories=None, exact_tones=False):
    """
    Filters words based on various criteria.
    - tones: List of tones to match (e.g., [1, 1]). Matches contiguous sequences.
    - categories: List of categories (any match)
    - exact_tones: If True, word_tones must exactly equal tones.
    """
    filtered = []
    for img_name, info in data.items():
        word_tones = info.get('tones', [])
        word_cats = info.get('categories', [])
        
        match = True
        
        if tones:
            if exact_tones:
                if word_tones != tones:
                    match = False
            else:
                if not has_sublist(word_tones, tones):
                    match = False
        
        # Filter by categories
        if match and categories:
            if not any(cat in word_cats for cat in categories):
                match = False
                
        if match:
            filtered.append((img_name, info))
            
    return filtered

def generate_doc(filtered_words, output_path, title, images_dir):
    doc = Document()
    doc.add_heading(title, 0)
    
    table = doc.add_table(rows=0, cols=2)
    
    for i in range(0, len(filtered_words), 2):
        row_cells = table.add_row().cells
        batch = filtered_words[i:i+2]
        
        for idx, (img_name, info) in enumerate(batch):
            cell = row_cells[idx]
            p = cell.add_paragraph()
            p.add_run(f"{info['hanzi']} ({info['pinyin']})\n").bold = True
            p.add_run(f"English: {info['english']}")
            
            img_path = os.path.join(images_dir, img_name)
            if os.path.exists(img_path):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2.5))
            else:
                cell.add_paragraph(f"[Image not found: {img_name}]")

    doc.save(output_path)
    print(f"Generated document with {len(filtered_words)} words at: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Generate Word documents from Chinese word data.")
    parser.add_argument("--tones", type=str, help="Comma separated tones to filter by (e.g. 1,1 for two first tones)")
    parser.add_argument("--exact-tones", action="store_true", help="If set, only words with PRECISELY the specified tones (counts and types) will be included.")
    parser.add_argument("--categories", type=str, help="Comma separated categories to filter by")
    parser.add_argument("--output", type=str, default="chinese_words_report.docx", help="Output filename")
    parser.add_argument("--title", type=str, default="Chinese Words Report", help="Document title")
    
    args = parser.parse_args()
    
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../'))
    words_data_path = os.path.join(base_dir, 'magic/words_data.js')
    images_dir = os.path.join(base_dir, 'magic/images')
    
    data = parse_js_data(words_data_path)
    
    tones_filter = [int(t) for t in args.tones.split(',') if t.strip()] if args.tones else None
    cats_filter = [c.strip() for c in args.categories.split(',') if c.strip()] if args.categories else None
    
    filtered = filter_words(data, tones=tones_filter, categories=cats_filter, exact_tones=args.exact_tones)
    
    generate_doc(filtered, args.output, args.title, images_dir)

if __name__ == "__main__":
    main()
